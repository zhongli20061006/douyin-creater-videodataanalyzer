"""Generate project technical documentation Word file."""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

PROJECT_ROOT = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(PROJECT_ROOT, '技术文档_抖音爬虫管理系统.docx')
os.environ.setdefault('SCRAPY_SETTINGS_MODULE', 'douyin_spider.settings')

doc = Document()

# ── Style setup ──
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.5

for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Microsoft YaHei'
    if level == 1:
        hs.font.size = Pt(18)
    elif level == 2:
        hs.font.size = Pt(14)
    elif level == 3:
        hs.font.size = Pt(12)

code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Consolas'
code_style.font.size = Pt(8)
code_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
code_style.paragraph_format.space_before = Pt(2)
code_style.paragraph_format.space_after = Pt(2)


def h(text, level=1):
    return doc.add_heading(text, level=level)


def p(text=None):
    return doc.add_paragraph(text or '')


def bold_run(para, text):
    run = para.add_run(text)
    run.bold = True
    return run


def code(lines):
    for line in lines:
        doc.add_paragraph(line, style='CodeBlock')


def table(headers, rows):
    ncols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=ncols)
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hdr in enumerate(headers):
        t.rows[0].cells[i].text = hdr
        for r in t.rows[0].cells[i].paragraphs:
            for run in r.runs:
                run.bold = True
                run.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
            for r in t.rows[ri + 1].cells[ci].paragraphs:
                for run in r.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()
    return t


# =============================================
# 封面
# =============================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('抖音爬虫管理系统\n技术文档')
title_run.font.size = Pt(28)
title_run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run('Scrapy + Redis + MySQL + FastAPI').font.size = Pt(14)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('版本：2.0.0\n日期：2026-06-08\n环境：Windows 11 + Python 3.13').font.size = Pt(11)

doc.add_page_break()

# =============================================
# 目录页 (手动)
# =============================================
h('目录', 1)
toc_items = [
    '一、项目概述',
    '二、技术栈与依赖',
    '三、项目结构',
    '四、系统架构',
    '五、核心模块详解',
    '    5.1   数据模型 (items.py)',
    '    5.2   爬虫引擎 (douyin_video.py)',
    '    5.3   数据管道 (pipelines.py)',
    '    5.4   中间件层 (middlewares.py)',
    '    5.5   配置中心 (settings.py)',
    '六、API 接口设计',
    '七、前端界面',
    '八、爬虫控制集成',
    '九、数据流与工作流程',
    '十、踩坑记录与疑难解决',
    '十一、部署与运维',
    '十二、性能与安全',
    '十三、项目成果与统计',
]
for item in toc_items:
    p(item)
doc.add_page_break()

# =============================================
# 一、项目概述
# =============================================
h('一、项目概述', 1)
p('本项目是一个完整的抖音视频数据采集与管理系统，使用 Scrapy 框架爬取抖音（douyin.com）视频数据，通过 Redis 分布式队列实现任务调度，使用 MySQL 持久化存储采集结果，并提供 FastAPI REST API 和 Web 管理面板用于数据查询、爬虫控制。')

h('核心功能', 2)
p('1. 抖音视频信息采集：支持按视频 ID、用户主页、搜索结果三种模式爬取')
p('2. 反反爬对抗：Playwright 浏览器渲染 + 多 UA 轮换 + Cookie 伪造 + 请求延迟')
p('3. 分布式任务调度：基于 Redis List 的任务队列，支持批量推送')
p('4. 数据管理 API：RESTful 风格接口，支持分页查询、搜索、排序、删除')
p('5. Web 管理面板：暗色主题 Dashboard，实时统计、视频列表、爬虫控制')
p('6. 爬虫集成控制：面板内一键启停爬虫，5 秒轮询状态')

h('原始项目背景', 2)
p('• 项目目录名为 "DjangoProject" 但实际是一个纯 Scrapy 项目（不含 Django）')
p('• 最初只有命令行操作方式，无任何 API 层和前端')
p('• 原开发者通过终端直接执行 SQL 查询数据')

doc.add_page_break()

# =============================================
# 二、技术栈与依赖
# =============================================
h('二、技术栈与依赖', 1)

table(['分层', '技术', '版本', '用途'], [
    ['语言', 'Python', '3.13', '全部后端逻辑'],
    ['爬虫框架', 'Scrapy', '2.12.0', '核心爬虫引擎'],
    ['分布式扩展', 'scrapy-redis', '0.7.3', 'Redis 分布式队列集成'],
    ['数据库', 'MySQL (PyMySQL)', '8.x / 1.1.0', '数据持久化存储'],
    ['缓存/队列', 'Redis', '6.x / 5.0.1', '任务队列 + 去重'],
    ['浏览器自动化', 'Playwright', 'latest', 'JS 渲染页面拦截 API'],
    ['UA 伪装', 'fake-useragent', '1.4.0', '随机 User-Agent'],
    ['JSON 解析', 'jsonpath', '0.82.2', '复杂 JSON 路径提取'],
    ['数据处理', 'pandas', '2.2.0', '数据分析（备用）'],
    ['Web 框架', 'FastAPI', '0.115+', 'REST API 服务'],
    ['ASGI 服务器', 'uvicorn', '0.34+', 'API 服务运行'],
    ['数据库驱动', 'pymysql', '1.1.0', 'MySQL 连接'],
    ['Redis 客户端', 'redis', '5.0.1', 'Redis 连接'],
    ['文档生成', 'python-docx', '1.2.0', 'Word 文档生成'],
])

h('MySQL 数据库结构', 2)
table(['字段', '类型', '说明'], [
    ['video_id', 'VARCHAR(64) PK', '抖音视频唯一ID'],
    ['video_title', 'VARCHAR(512)', '视频标题/描述'],
    ['video_desc', 'TEXT', '完整视频描述'],
    ['author_name', 'VARCHAR(128)', '作者昵称'],
    ['author_id', 'VARCHAR(64)', '作者唯一UID'],
    ['publish_time', 'DATETIME', '视频发布时间'],
    ['like_count', 'BIGINT', '点赞数'],
    ['comment_count', 'BIGINT', '评论数'],
    ['share_count', 'BIGINT', '分享数'],
    ['play_count', 'BIGINT', '播放量'],
    ['video_url', 'VARCHAR(1024)', '视频直链'],
    ['cover_url', 'VARCHAR(1024)', '封面图链接'],
    ['crawl_time', 'DATETIME', '爬取时间'],
    ['update_time', 'DATETIME', '最后更新时间'],
])

# =============================================
# 三、项目结构
# =============================================
doc.add_page_break()
h('三、项目结构', 1)

code([
    'PythonProject11/',
    '├── api.py                    # FastAPI 后端 (REST API + SpiderManager)',
    '├── run_backend.ps1           # 后台启动脚本 (pythonw.exe 无窗口)',
    '├── stop_backend.ps1          # 停止脚本 (读 PID 杀进程)',
    '├── start_spider.py           # [原始] CLI 启动爬虫 (Redis 队列驱动)',
    '├── batch_push.py             # [原始] 批量推送视频 ID 到 Redis',
    '├── id.py                     # [原始] 从用户主页提取视频 ID',
    '├── requirements.txt          # Python 依赖',
    '├── scrapy.cfg                # Scrapy 项目配置',
    '├── .gitignore                # Git 忽略规则',
    '├── video_ids.txt             # 待爬取视频 ID 列表',
    '│',
    '├── douyin_spider/            # Scrapy 核心包',
    '│   ├── settings.py           # 全局配置 (MySQL/Redis/Cookie/并发)',
    '│   ├── items.py              # 数据模型 (DouyinVideoItem)',
    '│   ├── pipelines.py          # MySQL 存储管道 (主)',
    '│   ├── middlewares.py        # 反反爬中间件 (UA/Cookie/延迟)',
    '│   ├── spiders/',
    '│   │   └── douyin_video.py   # 爬虫主逻辑 (视频/用户/搜索模式)',
    '│   └── utils/',
    '│       ├── pipelines.py      # [备用] PyMySQL 版管道',
    '│       └── db_helper.py      # 通用 MySQL 查询工具',
    '│',
    '├── frontend/                 # Web 管理面板',
    '│   ├── index.html            # 主页面 (Dashboard)',
    '│   ├── style.css             # 暗色主题样式',
    '│   └── app.js                # 前端逻辑 (API 调用/渲染/控制)',
    '│',
    '├── .idea/                    # PyCharm 配置',
    '│   └── runConfigurations/',
    '│       ├── 管理面板_后台启动.xml    # PowerShell 后台运行',
    '│       └── 管理面板_调试模式.xml    # Python 终端调试',
    '│',
    '└── .venv/                    # 虚拟环境 (不提交到 Git)',
])

# =============================================
# 四、系统架构
# =============================================
doc.add_page_break()
h('四、系统架构', 1)

p('整体架构采用分层设计：前端 -> API 层 -> 数据层 -> 爬虫层，各层通过明确接口通信。')
p()
table(['层级', '组件', '通信方式', '职责'], [
    ['表示层', 'frontend/ (HTML+CSS+JS)', 'HTTP (fetch API)', 'Web 管理面板，数据展示与操作'],
    ['API 层', 'api.py (FastAPI + uvicorn)', 'REST JSON', '数据 CRUD、爬虫控制、统计查询'],
    ['队列层', 'Redis List: douyin:start_urls', 'LPUSH / LRANGE', '爬虫任务分发与缓冲'],
    ['存储层', 'MySQL: douyin_spider.video_info', 'PyMySQL', '视频数据持久化与查询'],
    ['爬虫层', 'douyin_spider (Scrapy+Playwright)', 'Scrapy Engine', '页面抓取、反反爬、数据提取'],
    ['控制层', 'start_spider.py / batch_push.py', 'subprocess', '命令行任务推送与爬虫启动'],
])

h('端口与服务', 2)
table(['服务', '端口', '启动方式', '说明'], [
    ['MySQL', '3307', '系统服务', '非标准端口，root:见 local_config.py'],
    ['Redis', '6379', 'redis-server', '默认端口，存储任务队列'],
    ['FastAPI', '8001', 'run_backend.ps1 / PyCharm', '管理面板 + API'],
    ['Playwright', '--', 'Scrapy 内嵌', '非 headless，渲染 douyin.com'],
])

# =============================================
# 五、核心模块详解
# =============================================
doc.add_page_break()
h('五、核心模块详解', 1)

# 5.1
h('5.1 数据模型 — items.py', 2)
p('定义了 DouyinVideoItem，包含 13 个字段，对应 MySQL video_info 表。使用 scrapy.Field() 声明字段类型，爬虫 yield item 后自动流转到 Pipeline 处理。')
code([
    'class DouyinVideoItem(scrapy.Item):',
    '    video_id = scrapy.Field()       # 唯一标识 (primary key)',
    '    video_title = scrapy.Field()    # 视频标题',
    '    video_desc = scrapy.Field()     # 完整描述',
    '    author_name = scrapy.Field()    # 作者昵称',
    '    author_id = scrapy.Field()      # 作者 UID',
    '    publish_time = scrapy.Field()   # 发布时间的 datetime',
    '    like_count = scrapy.Field()     # 点赞数 (int)',
    '    comment_count = scrapy.Field()  # 评论数',
    '    share_count = scrapy.Field()    # 分享数',
    '    play_count = scrapy.Field()     # 播放量',
    '    video_url = scrapy.Field()      # 视频直链',
    '    cover_url = scrapy.Field()      # 封面图链接',
    '    crawl_time = scrapy.Field()     # 本次爬取时间',
])

# 5.2
h('5.2 爬虫引擎 — spiders/douyin_video.py', 2)
p('核心爬虫逻辑，支持三种爬取模式，从 Redis 队列消费任务：')

table(['模式', '任务类型', '数据来源', '采集范围'], [
    ['video', '单视频', 'Playwright 渲染 → 拦截 API / 解析 HTML', '单个视频全量信息'],
    ['user', '用户主页', 'aweme/v1/web/aweme/post/ API', '用户所有公开视频'],
    ['search', '搜索结果', 'jsonpath: $..aweme_info', '搜索结果视频列表'],
])

h('video 模式采集策略', 3)
p('1. Playwright 打开目标页面 → 拦截 API 请求 → 直接获取 JSON 响应（最优）')
p('2. 若拦截失败 → 解析 <script id="RENDER_DATA"> 中的 JSON（备用）')
p('3. 若以上均失败 → 用 CSS 选择器从 DOM 提取数据（兜底）')

h('关键 API 端点', 3)
code([
    '# 视频详情 API',
    'GET https://www.douyin.com/aweme/v1/web/aweme/detail/',
    '    ?aweme_id={video_id}&aid=1128',
    '# 用户主页 API',
    'GET https://www.douyin.com/aweme/v1/web/aweme/post/',
    '    ?sec_user_id={uid}&max_cursor={cursor}&count=20',
])

# 5.3
h('5.3 数据管道 — pipelines.py', 2)
p('MySQLPipeline 负责将 Item 写入 MySQL。关键设计：')

p('• 使用 INSERT ... ON DUPLICATE KEY UPDATE 实现 Upsert')
p('• 以 video_id 作为冲突键，已存在则更新所有动态字段')
p('• 更新字段包括：标题、描述、作者信息、互动数据、视频链接')
p('• 每次更新自动设置 update_time = NOW()')
p('• 连接在 spider 级别缓存，打开/关闭各执行一次')
p()
p('(注：utils/pipelines.py 是备用版本，使用 pymysql 驱动，仅更新 4 个计数字段，未被 settings.py 引用)')

# 5.4
h('5.4 中间件层 — middlewares.py', 2)
table(['优先级', '中间件', '策略'], [
    ['400', 'RandomUserAgentMiddleware', '轮换 3 个预设移动端 UA + fake_useragent 随机生成'],
    ['500', 'DouyinDownloaderMiddleware', '注入 Cookie、Referer、Origin 等抖音特定请求头'],
    ['543', 'PlaywrightMiddleware', '启动非 headless Chromium，禁用 AutomationControlled 标识，拦截 API 响应'],
    ['550', 'RequestDelayMiddleware', '请求间隔 3 秒，模拟人类浏览'],
    ['600', 'RetryMiddleware (内置)', 'HTTP 5xx + 408/429 自动重试，最多 3 次'],
])

# 5.5
h('5.5 配置中心 — settings.py', 2)
code([
    '# Redis',
    'REDIS_HOST = "localhost"      REDIS_PORT = 6379',
    'REDIS_START_URLS_KEY = "douyin:start_urls"',
    '',
    '# MySQL',
    'MYSQL_HOST = "localhost"      MYSQL_PORT = 3307',
    'MYSQL_USER = "root"           MYSQL_DB = "douyin_spider"',
    '',
    '# 并发控制',
    'CONCURRENT_REQUESTS = 1        # 单并发，降低检测',
    'DOWNLOAD_DELAY = 3             # 3秒间隔',
    '',
    '# Playwright',
    'TWISTED_REACTOR = ...AsyncioSelectorReactor',
    'DOWNLOAD_HANDLERS = { "https": PlaywrightDownloadHandler }',
    '',
    '# Cookie (需定期更新)',
    'DOUYIN_COOKIE = "..."',
])

doc.add_page_break()

# =============================================
# 六、API 接口设计
# =============================================
h('六、API 接口设计', 1)

h('数据查询类', 2)
table(['方法', '路径', '参数', '说明'], [
    ['GET', '/api/videos', 'page, page_size, search, sort_by, order', '分页查询视频列表（支持标题/作者/ID搜索）'],
    ['GET', '/api/videos/{video_id}', '—', '查询单个视频完整详情'],
    ['GET', '/api/stats', '—', '全量统计：总数/作者数/总互动/最新爬取时间/队列长度'],
])

h('爬取任务类', 2)
table(['方法', '路径', '参数', '说明'], [
    ['POST', '/api/crawl', '{"video_ids":[...],"task_type":"video"}', '推送视频 ID 到 Redis 任务队列 (LPUSH)'],
    ['DELETE', '/api/videos/{video_id}', '—', '从数据库删除指定视频记录'],
    ['GET', '/api/queue/length', '—', '查询 Redis 队列中的等待任务数量'],
])

h('爬虫控制类 (v2.0 新增)', 2)
table(['方法', '路径', '参数', '说明'], [
    ['POST', '/api/spider/start', '—', '启动 Scrapy 爬虫子进程 (subprocess.Popen)'],
    ['POST', '/api/spider/stop', '—', '终止爬虫子进程 (terminate → kill)'],
    ['GET', '/api/spider/status', '—', '返回 {running, pid, started_at}'],
    ['GET', '/api/spider/log', 'lines=50', '返回最近 N 行爬虫输出日志'],
])

h('数据格式', 2)
p('所有响应均为 JSON 格式，使用 Pydantic Models 定义 schema：')
code([
    'class VideoItem(BaseModel):       # 单个视频',
    'class PaginatedResponse(BaseModel): # 分页响应 {total, page, data[]}',
    'class StatsResponse(BaseModel):    # 统计响应',
    'class CrawlRequest(BaseModel):     # 爬取请求 {video_ids, task_type}',
])

doc.add_page_break()

# =============================================
# 七、前端界面
# =============================================
h('七、前端界面', 1)
p('纯 vanilla 实现，无前端框架依赖。单个 HTML + CSS + JS 文件，由 FastAPI StaticFiles 服务。')

h('组件清单', 2)
table(['组件', '技术', '功能'], [
    ['Header', 'CSS Grid', '标题、Redis 队列长度 badge、刷新按钮'],
    ['Stats Grid', 'CSS Grid auto-fit', '6 张统计卡片：视频总数/作者数/总赞/总评/总分/总播'],
    ['Spider Panel', 'Flexbox', '爬虫状态指示灯 + 启动/停止按钮（v2.0）'],
    ['Search Bar', 'Flexbox + input', '按视频ID/标题/作者搜索，支持回车触发'],
    ['Sort Controls', 'select', '按 6 个字段排序（爬取时间/点赞/评论/分享/播放/发布时间）'],
    ['Crawl Panel', 'toggle', '输入视频ID（文本/文件）→ 推送到 Redis 队列'],
    ['Data Table', 'table', '视频列表，hover 高亮，响应式'],
    ['Pagination', 'button row', '首页/上一页/页码/下一页/末页'],
    ['Detail Modal', 'overlay + grid', '视频 13 个字段完整展示'],
    ['Toast', 'fixed + animation', '成功/错误/信息提示，3 秒自动消失'],
    ['Loading', 'CSS animation', '旋转边框 spinner'],
])

h('设计规范', 2)
p('• 暗色主题（#0f1117 背景），CSS 自定义属性管理颜色体系')
p('• 响应式布局：768px 断点，2列统计卡片，全宽搜索栏')
p('• 安全处理：Douyin 资源链接仅展示为纯文本，不渲染 <img> 或 <a href>（防止 GFW 错误）')
p('• 数值格式化：>=10000 显示为 "x.x万"，>=1000 显示为 "x.xk"')

doc.add_page_break()

# =============================================
# 八、爬虫控制集成
# =============================================
h('八、爬虫控制集成', 1)
p('v2.0 核心新增功能，将原本独立的命令行操作集成到 Web 面板中。')

h('SpiderManager 类设计', 2)
code([
    'class SpiderManager:',
    '    def __init__(self):',
    '        self.process: subprocess.Popen | None',
    '        self.started_at: datetime | None',
    '        self.project_root / venv_python / log_path',
    '',
    '    def start() -> (bool, msg):',
    '        # 检查是否已在运行 → Popen 启动 start_spider.py',
    '        # stdout/stderr 重定向到 spider_output.log',
    '',
    '    def stop() -> (bool, msg):',
    '        # terminate() → wait(10s) → kill() 兜底',
    '',
    '    def is_alive() -> bool:',
    '        # self.process.poll() is None',
    '',
    '    def get_status() -> dict:',
    '        # {running, pid, started_at}',
    '',
    '    def get_log(lines=50) -> list[str]:',
    '        # 读取 spider_output.log 最后 N 行',
])

h('前端轮询机制', 2)
p('• 页面加载后立即查询一次爬虫状态')
p('• 每 5 秒通过 GET /api/spider/status 轮询')
p('• 根据 running 状态切换指示灯颜色（绿色=运行，灰色=停止）')
p('• 显示 PID 和运行时长（x 分钟前）')

p()
p('注意：爬虫只会在 API 进程内启动一次（单例），关闭 API 服务时子进程会一并终止。如需清理孤儿进程，需手动 kill。原始 CLI 方式完全保留，可在终端独立使用。')

doc.add_page_break()

# =============================================
# 十、踩坑记录与疑难解决
# =============================================
h('九、踩坑记录与疑难解决', 1)

bugs = [
    ('1. 端口占用冲突',
     '第一次启动时端口 8000 已被占用（可能是之前的进程未正确关闭）。切换至 8001 端口，并在每次启动前先检查端口占用情况。',
     '端口改为 8001 + 启动前 kill 旧进程'),
    ('2. 前端加载 Douyin 资源导致 GFW 拦截',
     '详情弹窗中使用了 <img src="douyin.com/..."> 和 <a href="douyin.com/..."> 直接加载 Douyin 资源。由于 GFW 封锁，浏览器长时间卡在 "正在连接..." 状态，页面变为空白色。',
     '所有 Douyin 链接改为纯文本展示，不渲染为可点击链接或图片。'),
    ('3. Emoji 编码错误 (GBK)',
     'start_server.py 中使用了 emoji 字符（如火箭图标），Windows 下 GBK 编码写入 .bat 文件时报错 UnicodeEncodeError。',
     '替换所有 emoji 为纯文本标记（如 [OK] 代替火箭符号）。'),
    ('4. Redis 模块未安装',
     '前端 Crawl Panel 调用 POST /api/crawl 时报错 ModuleNotFoundError: No module named redis。',
     'pip install redis，并将其加入 requirements.txt。'),
    ('5. MySQL 连接异常 (pymysql.err)',
     'get_db() 原使用 yield 生成器模式，在某些并发场景下导致连接被提前关闭（Double-close 错误）。',
     '改为直接 pymysql.connect() 返回，在调用处 finally 中显式 db_close()，增加 connect_timeout=5 参数。'),
    ('6. FastAPI Query 参数 deprecated',
     '使用了已弃用的 regex 参数（Query(regex=...)），FastAPI 新版本只支持 pattern。',
     '将 Query(regex=...) 替换为 Query(pattern=...)。'),
    ('7. Playwright 浏览器未安装',
     '爬虫启动后报错：chromium-1208/chrome-win64/chrome.exe 不存在。原因是 pip install playwright 只安装了 Python 包，未安装浏览器二进制文件。且本项目的 playwright 包版本与其他项目不一致，共享缓存目录 C:\\PlaywrightBrowsers\\ 导致版本冲突。',
     'playwright install chromium，升级 playwright 包解决版本不一致。'),
    ('8. Playwright 安装镜像失败',
     'playwright install chromium 走 npmmirror 镜像时返回 404（该版本未镜像缓存）。',
     '设置 PLAYWRIGHT_DOWNLOAD_HOST 为官方 CDN 直连下载。'),
    ('9. Uvicorn 进程僵死',
     '前端持续 "连接中..." 状态，curl 请求超时。原因是 uvicorn 进程虽然存活但已进入僵死状态（无法处理请求）。',
     'Stop-Process -Force 强杀 → 重新启动。改用 pythonw.exe + subprocess.Popen 的独立进程管理方案。'),
    ('10. .gitignore 误排除 frontend/index.html',
     '.gitignore 中 *.html 规则原意是排除爬虫抓取的 HTML 快照（douyin_page.html），但同时也排除了前端需要的 index.html。',
     '添加 !frontend/*.html 例外规则，使前端 HTML 可以被 Git 追踪。'),
    ('11. PowerShell 后台进程问题',
     'Start-Process -WindowStyle Hidden 配合 pythonw.exe 时，进程启动后立即退出，端口未监听。',
     '改用 ProcessStartInfo + UseShellExecute=true + WindowStyle=Hidden 方式创建真正独立的后台进程。'),
]

for title, desc, fix in bugs:
    h(title, 2)
    p(f'问题描述：{desc}')
    p(f'解决方案：{fix}')
    doc.add_paragraph()

doc.add_page_break()

# =============================================
# 十一、部署与运维
# =============================================
h('十、部署与运维', 1)

h('环境要求', 2)
p('• Windows 10+ / Windows Server 2019+')
p('• Python 3.13（虚拟环境 .venv）')
p('• MySQL 8.x（端口 3307）')
p('• Redis 6.x+（端口 6379）')
p('• Playwright Chromium 浏览器（C:\\PlaywrightBrowsers\\）')

h('启动流程', 2)
p('1. 启动 MySQL 服务（确保 3307 端口可访问）')
p('2. 启动 Redis: redis-server')
p('3. 安装依赖: pip install -r requirements.txt')
p('4. 安装浏览器: playwright install chromium')
p('5. 启动管理面板: 任选一种方式')
p()

table(['方式', '命令/操作', '特点'], [
    ['后台无窗口', '.\\run_backend.ps1', 'pythonw.exe 静默运行，浏览器关闭不影响'],
    ['终端调试', 'uvicorn api:app --port 8001', '查看实时日志，Ctrl+C 停止'],
    ['PyCharm', '运行 "管理面板_后台启动" 或 "管理面板_调试模式"', 'IDE 集成，支持断点调试'],
])

h('停止服务', 2)
p('• 停止 API: .\\stop_backend.ps1')
p('• 停止爬虫: 前端面板点击 "停止爬虫" 或 kill 对应 python 进程')

h('重启与数据恢复', 2)
p('• 管理面板重启：重新执行 run_backend.ps1 即可')
p('• 爬虫故障恢复：Redis 队列数据持久化，重启爬虫后自动继续消费')
p('• MySQL 数据：INSERT ... ON DUPLICATE KEY UPDATE 保证幂等，重复爬取自动更新')

doc.add_page_break()

# =============================================
# 十二、性能与安全
# =============================================
h('十一、性能与安全', 1)

h('性能设计', 2)
p('• 单并发爬取（CONCURRENT_REQUESTS=1），避免触发反爬阈值')
p('• 请求间隔 3 秒（DOWNLOAD_DELAY=3），模拟人类浏览行为')
p('• 前端无框架依赖，页面加载体积 < 50KB（Gzip 后约 15KB）')
p('• API 分页查询，默认每页 20 条，避免全量数据传输')
p('• MySQL 使用 DictCursor 直接返回字典格式，减少序列化开销')

h('安全注意事项', 2)
p('• MySQL 密码硬编码在 settings.py 和 api.py 的回退配置中 → 生产环境应使用环境变量')
p('• Douyin Cookie 硬编码在 settings.py → 需定期更新（约 2-4 小时过期）')
p('• CORS 放开至 *（开发环境）→ 生产环境应限制为具体域名')
p('• 无用户认证机制 → 仅限内网使用，不宜暴露到公网')
p('• video_ids.txt 包含真实抖音 ID → 已在 .gitignore 中排除')

doc.add_page_break()

# =============================================
# 十三、项目成果
# =============================================
h('十二、项目成果与统计', 1)
p('截至文档编写时（2026-06-08），项目状态：')

table(['指标', '数值'], [
    ['总采集视频数', '182'],
    ['覆盖作者数', '5'],
    ['总点赞数', '6,743,375（674万）'],
    ['总评论数', '413,271（41万）'],
    ['总分享数', '451,004（45万）'],
    ['Redis 队列待处理', '0（已全部消费）'],
    ['最新采集时间', '2026-05-20 11:36'],
    ['Scrapy 爬虫项数', '1（douyin_video）'],
    ['API 端点数', '8'],
    ['前端组件数', '11'],
    ['Git 提交数', '3'],
])

# ── Save ──
doc.save(OUTPUT)
print(f'Document saved: {OUTPUT}')
